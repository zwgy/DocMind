"""内置文档导出 MCP。"""

from __future__ import annotations

import tempfile
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from mcp.server.fastmcp import FastMCP
from openpyxl import Workbook

from yuxi.services.file_preview import convert_office_to_pdf

_ARTIFACT_DIR = Path(tempfile.gettempdir()) / "yuxi-mcp-artifacts"
mcp = FastMCP("document-exporter")


def _filename(filename: str, suffix: str) -> str:
    name = Path(filename or "document").name
    return f"{Path(name).stem or 'document'}{suffix}"


def _store(filename: str, content: bytes) -> dict[str, str]:
    _ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    path = _ARTIFACT_DIR / f"{uuid.uuid4().hex}{Path(filename).suffix}"
    path.write_bytes(content)
    return {"artifact_path": str(path), "filename": filename}


def _docx(content: str, title: str | None = None) -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Microsoft YaHei"
    document.styles["Normal"].font.size = Pt(10.5)
    if title:
        document.add_heading(title, level=0)
    for line in content.splitlines() or [""]:
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        else:
            document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@mcp.tool(description="根据文本生成 DOCX 文件。文件会自动写入当前对话的 outputs 目录。")
def generate_docx(filename: str, content: str, title: str | None = None) -> dict[str, str]:
    return _store(_filename(filename, ".docx"), _docx(content, title))


@mcp.tool(description="根据文本生成 PDF 文件。文件会自动写入当前对话的 outputs 目录。")
async def generate_pdf(filename: str, content: str, title: str | None = None) -> dict[str, str]:
    pdf = await convert_office_to_pdf("document.docx", _docx(content, title))
    return _store(_filename(filename, ".pdf"), pdf)


@mcp.tool(description="根据工作表数据生成 XLSX 文件。文件会自动写入当前对话的 outputs 目录。")
def generate_xlsx(filename: str, sheets: dict[str, list[list[Any]]]) -> dict[str, str]:
    if not sheets:
        raise ValueError("至少需要提供一个工作表")
    workbook = Workbook()
    workbook.remove(workbook.active)
    for sheet_name, rows in sheets.items():
        worksheet = workbook.create_sheet(title=sheet_name)
        for row in rows:
            worksheet.append(row)
    output = BytesIO()
    workbook.save(output)
    return _store(_filename(filename, ".xlsx"), output.getvalue())


if __name__ == "__main__":
    mcp.run(transport="stdio")

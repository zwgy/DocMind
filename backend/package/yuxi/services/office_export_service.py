"""受限 Office 定义的离线导出服务。"""

from __future__ import annotations

import asyncio
import json
import os
import re
import subprocess
import tempfile
import uuid
from collections.abc import Callable
from io import BytesIO
from pathlib import Path
from typing import Annotated, Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.drawing.image import Image as WorksheetImage
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image
from pydantic import BaseModel, ConfigDict, Field, TypeAdapter, field_validator, model_validator

from yuxi.services.file_preview import convert_office_to_pdf

_MAX_DEFINITION_BYTES = 2 * 1024 * 1024
_MAX_IMAGE_PIXELS = 100_000_000
_MAX_SHEETS = 20
_MAX_ROWS_PER_SHEET = 20_000
_DOCX_IMAGE_VERTICAL_RESERVE_CM = 3.5
_INVALID_FILENAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
_CELL_ANCHOR_RE = re.compile(r"^[A-Z]{1,3}[1-9][0-9]{0,6}$")
_SHEET_INVALID_RE = re.compile(r"[\[\]:*?/\\]")
_SUPPORTED_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp", ".svg"}

type CellValue = str | int | float | bool | None
type SourceResolver = Callable[[str, set[str]], Path]


class OfficeExportError(ValueError):
    """向模型返回的可执行中文导出错误。"""


class _StrictModel(BaseModel):
    model_config = ConfigDict(extra="forbid")


class HeadingBlock(_StrictModel):
    type: Literal["heading"]
    text: str = Field(min_length=1, max_length=500)
    level: int = Field(default=1, ge=1, le=4)


class ParagraphBlock(_StrictModel):
    type: Literal["paragraph"]
    text: str = Field(max_length=20_000)


class TableBlock(_StrictModel):
    type: Literal["table"]
    rows: list[list[CellValue]] = Field(min_length=1, max_length=2_000)
    header: bool = True

    @model_validator(mode="after")
    def validate_rows(self) -> TableBlock:
        width = len(self.rows[0])
        if width == 0 or width > 30:
            raise ValueError("表格列数必须在 1 到 30 之间")
        if any(len(row) != width for row in self.rows):
            raise ValueError("表格每一行的列数必须一致")
        return self


class ImageBlock(_StrictModel):
    type: Literal["image"]
    source_path: str = Field(min_length=1, max_length=500)
    caption: str | None = Field(default=None, max_length=500)
    width_cm: float = Field(default=16, ge=2, le=17)


class PageBreakBlock(_StrictModel):
    type: Literal["page_break"]


DocumentBlock = Annotated[
    HeadingBlock | ParagraphBlock | TableBlock | ImageBlock | PageBreakBlock,
    Field(discriminator="type"),
]


class DocumentDefinition(_StrictModel):
    kind: Literal["document"]
    title: str | None = Field(default=None, min_length=1, max_length=500)
    blocks: list[DocumentBlock] = Field(min_length=1, max_length=500)


class WorksheetImageDefinition(_StrictModel):
    source_path: str = Field(min_length=1, max_length=500)
    anchor: str = Field(pattern=_CELL_ANCHOR_RE.pattern)
    width_px: int = Field(default=960, ge=64, le=2400)


class WorksheetDefinition(_StrictModel):
    name: str = Field(min_length=1, max_length=31)
    rows: list[list[CellValue]] = Field(default_factory=list, max_length=_MAX_ROWS_PER_SHEET)
    header_rows: int = Field(default=1, ge=0, le=10)
    freeze_panes: str | None = Field(default=None, pattern=_CELL_ANCHOR_RE.pattern)
    images: list[WorksheetImageDefinition] = Field(default_factory=list, max_length=20)

    @field_validator("name")
    @classmethod
    def validate_name(cls, value: str) -> str:
        normalized = value.strip()
        if not normalized or _SHEET_INVALID_RE.search(normalized) or normalized.endswith("'"):
            raise ValueError("工作表名称包含 Excel 不允许的字符")
        return normalized

    @model_validator(mode="after")
    def validate_sheet(self) -> WorksheetDefinition:
        if not self.rows and not self.images:
            raise ValueError("工作表至少需要一行数据或一张图片")
        if self.rows and self.header_rows > len(self.rows):
            raise ValueError("header_rows 不能超过工作表数据行数")
        if any(len(row) > 100 for row in self.rows):
            raise ValueError("工作表单行不能超过 100 列")
        return self


class WorkbookDefinition(_StrictModel):
    kind: Literal["workbook"]
    sheets: list[WorksheetDefinition] = Field(min_length=1, max_length=_MAX_SHEETS)

    @model_validator(mode="after")
    def validate_sheets(self) -> WorkbookDefinition:
        names = [sheet.name.casefold() for sheet in self.sheets]
        if len(names) != len(set(names)):
            raise ValueError("工作表名称不能重复")
        return self


OfficeDefinition = Annotated[DocumentDefinition | WorkbookDefinition, Field(discriminator="kind")]
OFFICE_DEFINITION_ADAPTER = TypeAdapter(OfficeDefinition)


def load_office_definition(path: Path) -> DocumentDefinition | WorkbookDefinition:
    if path.suffix.lower() != ".json":
        raise OfficeExportError("definition_path 必须是 JSON 文件")
    if not path.is_file():
        raise OfficeExportError("Office 定义文件不存在")
    if path.stat().st_size > _MAX_DEFINITION_BYTES:
        raise OfficeExportError("Office 定义文件超过 2 MB 限制")
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise OfficeExportError("Office 定义文件不是有效的 UTF-8 JSON") from exc
    if not isinstance(payload, dict) or payload.get("kind") not in {"document", "workbook"}:
        raise OfficeExportError("Office 定义顶层字段 kind 必须为 document（DOCX/PDF）或 workbook（XLSX）")
    try:
        return OFFICE_DEFINITION_ADAPTER.validate_python(payload)
    except ValueError as exc:
        raise OfficeExportError(f"Office 定义不符合格式要求：{exc}") from exc


def validate_definition_format(
    definition: DocumentDefinition | WorkbookDefinition,
    output_format: Literal["docx", "pdf", "xlsx"],
) -> None:
    if output_format in {"docx", "pdf"} and definition.kind != "document":
        raise OfficeExportError("DOCX/PDF 必须使用 kind=document 的定义")
    if output_format == "xlsx" and definition.kind != "workbook":
        raise OfficeExportError("XLSX 必须使用 kind=workbook 的定义")


def _normalize_output_name(output_name: str) -> str:
    value = str(output_name or "").strip()
    if (
        not value
        or len(value) > 100
        or value in {".", ".."}
        or value.endswith((".", " "))
        or _INVALID_FILENAME_RE.search(value)
        or Path(value).suffix
    ):
        raise OfficeExportError("output_name 必须是不含扩展名和路径的 1 至 100 字符文件名")
    return value


def _reserve_output(directory: Path, output_name: str, suffix: str) -> tuple[Path, str]:
    directory.mkdir(parents=True, exist_ok=True)
    for index in range(1, 1000):
        numbered = "" if index == 1 else f"-{index}"
        filename = f"{output_name}{numbered}{suffix}"
        target = directory / filename
        try:
            descriptor = os.open(target, os.O_CREAT | os.O_EXCL | os.O_WRONLY, 0o600)
            os.close(descriptor)
            return target, filename
        except FileExistsError:
            continue
    raise OfficeExportError("同名导出文件过多，请更换 output_name")


class _ImageMaterializer:
    """把 Office 不稳定支持的图片格式收敛成临时 PNG。"""

    def __init__(self, resolver: SourceResolver, temporary_directory: Path):
        self._resolver = resolver
        self._temporary_directory = temporary_directory
        self._cache: dict[tuple[str, int], Path] = {}

    def materialize(self, source_path: str, *, width_px: int) -> Path:
        key = (source_path, width_px)
        cached = self._cache.get(key)
        if cached is not None:
            return cached

        source = self._resolver(source_path, _SUPPORTED_IMAGE_SUFFIXES)
        suffix = source.suffix.lower()
        if suffix == ".svg":
            target = self._temporary_directory / f"{uuid.uuid4().hex}.png"
            self._convert_svg(source, target, width_px)
        elif suffix in {".png", ".jpg", ".jpeg"}:
            self._validate_raster(source)
            target = source
        else:
            target = self._temporary_directory / f"{uuid.uuid4().hex}.png"
            self._convert_raster(source, target)
        self._cache[key] = target
        return target

    @staticmethod
    def _validate_raster(path: Path) -> None:
        try:
            with Image.open(path) as image:
                if image.width * image.height > _MAX_IMAGE_PIXELS:
                    raise OfficeExportError("图片像素总量超过 1 亿限制")
                image.verify()
        except OfficeExportError:
            raise
        except Exception as exc:
            raise OfficeExportError(f"图片无法读取：{path.name}") from exc

    @staticmethod
    def _convert_raster(source: Path, target: Path) -> None:
        try:
            with Image.open(source) as image:
                if image.width * image.height > _MAX_IMAGE_PIXELS:
                    raise OfficeExportError("图片像素总量超过 1 亿限制")
                frame = image.convert("RGBA" if image.mode in {"RGBA", "LA"} else "RGB")
                frame.save(target, format="PNG")
        except OfficeExportError:
            raise
        except Exception as exc:
            raise OfficeExportError(f"图片格式转换失败：{source.name}") from exc

    @staticmethod
    def _convert_svg(source: Path, target: Path, width_px: int) -> None:
        script = Path(__file__).resolve().parent / "scripts" / "svg_to_png.mjs"
        if not script.is_file():
            raise OfficeExportError("SVG 图片转换器未安装")
        payload = json.dumps(
            {"source": str(source), "output": str(target), "width": width_px},
            ensure_ascii=False,
        )
        try:
            result = subprocess.run(
                ["node", str(script)],
                input=payload,
                text=True,
                capture_output=True,
                timeout=30,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise OfficeExportError("SVG 图片转换器启动失败或超时") from exc
        if result.returncode != 0 or not target.is_file():
            detail = (result.stderr or result.stdout).strip().splitlines()
            raise OfficeExportError(f"SVG 图片转换失败：{(detail[-1] if detail else '转换器未生成文件')[:300]}")


def _set_run_font(run, *, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Noto Sans CJK SC"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _cell_text(value: CellValue) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    return str(value)


def _build_docx(definition: DocumentDefinition, resolver: SourceResolver) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    content_width_cm = (section.page_width - section.left_margin - section.right_margin) / Cm(1)
    content_height_cm = (section.page_height - section.top_margin - section.bottom_margin) / Cm(1)
    max_image_height_cm = content_height_cm - _DOCX_IMAGE_VERTICAL_RESERVE_CM

    normal = document.styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)

    with tempfile.TemporaryDirectory(prefix="yuxi-office-images-") as temporary:
        images = _ImageMaterializer(resolver, Path(temporary))
        if definition.title:
            title = document.add_heading(definition.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                _set_run_font(run, size=20, bold=True)

        pending_page_break = False
        for block in definition.blocks:
            if isinstance(block, PageBreakBlock):
                # 将分页要求附着到下一个内容块，避免高图恰好占满页面时，独立分页符
                # 被排到下一页后再次分页，从而在 LibreOffice/PDF 中产生整页空白。
                pending_page_break = True
                continue

            if isinstance(block, HeadingBlock):
                paragraph = document.add_heading(block.text, level=block.level)
                paragraph.paragraph_format.page_break_before = pending_page_break
                for run in paragraph.runs:
                    _set_run_font(run, bold=True)
            elif isinstance(block, ParagraphBlock):
                paragraph = document.add_paragraph(block.text)
                paragraph.paragraph_format.page_break_before = pending_page_break
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.line_spacing = 1.35
            elif isinstance(block, TableBlock):
                if pending_page_break:
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.page_break_before = True
                table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
                table.style = "Table Grid"
                for row_index, row in enumerate(block.rows):
                    for column_index, value in enumerate(row):
                        cell = table.cell(row_index, column_index)
                        cell.text = _cell_text(value)
                        if block.header and row_index == 0:
                            _shade_cell(cell, "E9EEF5")
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    _set_run_font(run, bold=True)
                document.add_paragraph()
            elif isinstance(block, ImageBlock):
                # Word 对 SVG 的支持取决于客户端版本；统一物化为高分辨率 PNG 才能保证
                # Word、WPS 和 LibreOffice 三端的文档内容一致。
                width_px = min(2400, max(600, round(block.width_cm / 2.54 * 240)))
                source = images.materialize(block.source_path, width_px=width_px)
                with Image.open(source) as image:
                    image_width_px, image_height_px = image.size

                # width_cm 是期望最大宽度。纵向流程图如果只按宽度写入会高出页面，
                # 因此还要为同页标题、图题和段落间距预留空间，再按原始比例整体缩小。
                picture_width_cm = min(block.width_cm, content_width_cm)
                picture_height_cm = picture_width_cm * image_height_px / image_width_px
                if picture_height_cm > max_image_height_cm:
                    scale = max_image_height_cm / picture_height_cm
                    picture_width_cm *= scale
                    picture_height_cm = max_image_height_cm

                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.page_break_before = pending_page_break
                paragraph.add_run().add_picture(
                    str(source),
                    width=Cm(picture_width_cm),
                    height=Cm(picture_height_cm),
                )
                if block.caption:
                    caption = document.add_paragraph(block.caption, style="Caption")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        _set_run_font(run, size=9)
            pending_page_break = False

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _write_excel_value(cell, value: CellValue) -> None:
    cell.value = value
    # 首版不开放公式能力。把等号开头内容强制写为字符串，可避免外部数据在打开
    # 工作簿时被 Excel 当作公式或外部链接执行。
    if isinstance(value, str) and value.startswith(("=", "+", "-", "@")):
        cell.data_type = "s"
    cell.alignment = Alignment(vertical="top", wrap_text=True)


def _build_xlsx(definition: WorkbookDefinition, resolver: SourceResolver) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    with tempfile.TemporaryDirectory(prefix="yuxi-office-images-") as temporary:
        images = _ImageMaterializer(resolver, Path(temporary))
        for sheet_definition in definition.sheets:
            worksheet = workbook.create_sheet(sheet_definition.name)
            dominant_image_size: tuple[int, int] | None = None
            for row_index, row in enumerate(sheet_definition.rows, start=1):
                for column_index, value in enumerate(row, start=1):
                    cell = worksheet.cell(row=row_index, column=column_index)
                    _write_excel_value(cell, value)
                    if row_index <= sheet_definition.header_rows:
                        cell.font = Font(name="Noto Sans CJK SC", bold=True, color="1F2937")
                        cell.fill = PatternFill("solid", fgColor="E9EEF5")
                    else:
                        cell.font = Font(name="Noto Sans CJK SC", color="1F2937")

            if sheet_definition.freeze_panes:
                worksheet.freeze_panes = sheet_definition.freeze_panes
            elif sheet_definition.header_rows and sheet_definition.rows:
                worksheet.freeze_panes = f"A{sheet_definition.header_rows + 1}"

            max_columns = max((len(row) for row in sheet_definition.rows), default=0)
            for column_index in range(1, max_columns + 1):
                values = [
                    _cell_text(row[column_index - 1]) for row in sheet_definition.rows[:500] if len(row) >= column_index
                ]
                width = min(45, max(10, max((len(value) for value in values), default=0) + 2))
                worksheet.column_dimensions[get_column_letter(column_index)].width = width

            for image_definition in sheet_definition.images:
                source = images.materialize(image_definition.source_path, width_px=image_definition.width_px)
                image = WorksheetImage(str(source))
                if image.width:
                    ratio = image_definition.width_px / image.width
                    image.width = image_definition.width_px
                    image.height = round(image.height * ratio)
                image_size = (int(image.width), int(image.height))
                if dominant_image_size is None or image_size[0] * image_size[1] > (
                    dominant_image_size[0] * dominant_image_size[1]
                ):
                    dominant_image_size = image_size
                worksheet.add_image(image, image_definition.anchor)

            if dominant_image_size is not None:
                # 图片在 Excel 编辑视图中是浮动绘图对象，不会被单元格裁切；这里补充打印
                # 语义，让 LibreOffice/PDF 预览也按图片主方向排版，避免宽图被横向截断。
                worksheet.sheet_view.showGridLines = False
                worksheet.sheet_properties.pageSetUpPr.fitToPage = True
                worksheet.page_setup.paperSize = worksheet.PAPERSIZE_A4
                worksheet.page_setup.orientation = (
                    worksheet.ORIENTATION_LANDSCAPE
                    if dominant_image_size[0] >= dominant_image_size[1]
                    else worksheet.ORIENTATION_PORTRAIT
                )
                worksheet.page_setup.fitToWidth = 1
                # 展示型工作表通常只有标题或少量说明，适合把整张静态图收进一页；
                # 数据量大的工作表保留纵向分页，避免为了图片把表格正文缩小到不可读。
                worksheet.page_setup.fitToHeight = 1 if len(sheet_definition.rows) <= 50 else 0
                worksheet.page_margins.left = 0.25
                worksheet.page_margins.right = 0.25
                worksheet.page_margins.top = 0.35
                worksheet.page_margins.bottom = 0.35
                worksheet.print_options.horizontalCentered = True

        # openpyxl 在 save() 时才读取图片内容，因此保存必须发生在 SVG 等临时 PNG
        # 仍处于生命周期内的阶段；提前退出临时目录会生成无法写入图片的工作簿。
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()


async def export_office_file(
    *,
    definition_path: Path,
    output_format: Literal["docx", "pdf", "xlsx"],
    output_name: str,
    output_directory: Path,
    virtual_output_directory: str,
    source_resolver: SourceResolver,
) -> dict[str, str]:
    definition = load_office_definition(definition_path)
    validate_definition_format(definition, output_format)
    normalized_name = _normalize_output_name(output_name)
    final, filename = _reserve_output(output_directory, normalized_name, f".{output_format}")
    temporary = output_directory / f".{uuid.uuid4().hex}.office.tmp"
    try:
        if output_format in {"docx", "pdf"}:
            assert isinstance(definition, DocumentDefinition)
            docx = await asyncio.to_thread(_build_docx, definition, source_resolver)
            content = docx if output_format == "docx" else await convert_office_to_pdf("document.docx", docx)
        else:
            assert isinstance(definition, WorkbookDefinition)
            content = await asyncio.to_thread(_build_xlsx, definition, source_resolver)
        temporary.write_bytes(content)
        os.replace(temporary, final)
    except OfficeExportError:
        raise
    except Exception as exc:
        raise OfficeExportError(f"{output_format.upper()} 导出失败：{exc}") from exc
    finally:
        temporary.unlink(missing_ok=True)
        if final.exists() and final.stat().st_size == 0:
            final.unlink(missing_ok=True)

    return {
        "artifact_path": f"{virtual_output_directory.rstrip('/')}/{filename}",
        "summary": f"已生成 {output_format.upper()} 文件：{filename}",
    }

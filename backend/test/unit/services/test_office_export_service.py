from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook
from PIL import Image

from yuxi.services import office_export_service as service


def _resolver(paths: dict[str, Path]):
    def resolve(virtual_path: str, suffixes: set[str]) -> Path:
        path = paths[virtual_path]
        assert path.suffix.lower() in suffixes
        return path

    return resolve


def _png(path: Path, size: tuple[int, int] = (320, 180)) -> None:
    Image.new("RGB", size, "#2f6f5e").save(path)


@pytest.mark.asyncio
async def test_export_docx_inserts_table_picture_and_caption(tmp_path: Path) -> None:
    image = tmp_path / "chart.png"
    definition = tmp_path / "report.json"
    outputs = tmp_path / "outputs"
    _png(image)
    definition.write_text(
        json.dumps(
            {
                "kind": "document",
                "title": "检查报告",
                "blocks": [
                    {"type": "heading", "level": 1, "text": "检查结论"},
                    {"type": "table", "rows": [["项目", "数量"], ["问题", 2]]},
                    {
                        "type": "image",
                        "source_path": "/home/gem/user-data/outputs/chart.png",
                        "caption": "图 1 风险分布",
                        "width_cm": 12,
                    },
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = await service.export_office_file(
        definition_path=definition,
        output_format="docx",
        output_name="检查报告",
        output_directory=outputs,
        virtual_output_directory="/home/gem/user-data/outputs",
        source_resolver=_resolver({"/home/gem/user-data/outputs/chart.png": image}),
    )

    document = Document(outputs / "检查报告.docx")
    assert result["artifact_path"].endswith("/检查报告.docx")
    assert [row.cells[0].text for row in document.tables[0].rows] == ["项目", "问题"]
    assert len(document.inline_shapes) == 1
    assert any(paragraph.text == "图 1 风险分布" for paragraph in document.paragraphs)
    assert {path.suffix for path in outputs.iterdir()} == {".docx"}


@pytest.mark.asyncio
async def test_export_docx_scales_tall_picture_within_page(tmp_path: Path) -> None:
    image = tmp_path / "flow.png"
    definition = tmp_path / "report.json"
    outputs = tmp_path / "outputs"
    _png(image, (400, 1600))
    definition.write_text(
        json.dumps(
            {
                "kind": "document",
                "blocks": [
                    {
                        "type": "image",
                        "source_path": "/home/gem/user-data/outputs/flow.png",
                        "caption": "图 1 纵向流程图",
                        "width_cm": 16,
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    await service.export_office_file(
        definition_path=definition,
        output_format="docx",
        output_name="纵向流程图",
        output_directory=outputs,
        virtual_output_directory="/home/gem/user-data/outputs",
        source_resolver=_resolver({"/home/gem/user-data/outputs/flow.png": image}),
    )

    document = Document(outputs / "纵向流程图.docx")
    shape = document.inline_shapes[0]
    section = document.sections[0]
    max_height_cm = (section.page_height - section.top_margin - section.bottom_margin) / service.Cm(
        1
    ) - service._DOCX_IMAGE_VERTICAL_RESERVE_CM
    assert shape.height / service.Cm(1) <= max_height_cm + 0.01
    assert shape.width / shape.height == pytest.approx(0.25, abs=0.001)


@pytest.mark.asyncio
async def test_export_xlsx_inserts_rows_and_picture(tmp_path: Path) -> None:
    image = tmp_path / "chart.svg"
    definition = tmp_path / "ledger.json"
    outputs = tmp_path / "outputs"
    image.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" width="320" height="180">'
        '<rect width="320" height="180" fill="#2f6f5e"/></svg>',
        encoding="utf-8",
    )
    definition.write_text(
        json.dumps(
            {
                "kind": "workbook",
                "sheets": [
                    {
                        "name": "统计",
                        "rows": [["类型", "数量"], ["风险", 2]],
                        "images": [
                            {
                                "source_path": "/home/gem/user-data/outputs/chart.svg",
                                "anchor": "D2",
                                "width_px": 640,
                            }
                        ],
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    await service.export_office_file(
        definition_path=definition,
        output_format="xlsx",
        output_name="风险台账",
        output_directory=outputs,
        virtual_output_directory="/home/gem/user-data/outputs",
        source_resolver=_resolver({"/home/gem/user-data/outputs/chart.svg": image}),
    )

    workbook = load_workbook(outputs / "风险台账.xlsx")
    worksheet = workbook["统计"]
    assert list(worksheet.values) == [("类型", "数量"), ("风险", 2)]
    assert worksheet.freeze_panes == "A2"
    assert len(worksheet._images) == 1


@pytest.mark.asyncio
async def test_export_pdf_reuses_docx_conversion(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    definition = tmp_path / "report.json"
    outputs = tmp_path / "outputs"
    definition.write_text(
        '{"kind":"document","blocks":[{"type":"paragraph","text":"正文"}]}',
        encoding="utf-8",
    )

    async def fake_convert(filename: str, content: bytes) -> bytes:
        assert filename == "document.docx"
        assert content.startswith(b"PK")
        return b"%PDF-1.7\n"

    monkeypatch.setattr(service, "convert_office_to_pdf", fake_convert)
    await service.export_office_file(
        definition_path=definition,
        output_format="pdf",
        output_name="报告",
        output_directory=outputs,
        virtual_output_directory="/home/gem/user-data/outputs",
        source_resolver=_resolver({}),
    )

    assert (outputs / "报告.pdf").read_bytes() == b"%PDF-1.7\n"


def test_svg_materializer_generates_png_only_on_demand(tmp_path: Path) -> None:
    svg = tmp_path / "chart.svg"
    temporary = tmp_path / "temporary"
    temporary.mkdir()
    svg.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" width="200" height="100">'
        '<rect width="200" height="100" fill="#2f6f5e"/></svg>',
        encoding="utf-8",
    )
    images = service._ImageMaterializer(
        _resolver({"/home/gem/user-data/outputs/chart.svg": svg}),
        temporary,
    )

    output = images.materialize("/home/gem/user-data/outputs/chart.svg", width_px=800)

    assert output.read_bytes().startswith(b"\x89PNG\r\n\x1a\n")
    assert svg.read_text(encoding="utf-8").startswith("<svg")


def test_definition_format_must_match_output_format() -> None:
    document = service.OFFICE_DEFINITION_ADAPTER.validate_python(
        {"kind": "document", "blocks": [{"type": "paragraph", "text": "正文"}]}
    )

    with pytest.raises(service.OfficeExportError, match="XLSX"):
        service.validate_definition_format(document, "xlsx")
